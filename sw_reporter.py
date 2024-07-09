import os
os.environ['TRANSFORMERS_NO_ADVISORY_WARNINGS'] = '1'
import warnings
warnings.filterwarnings("ignore", message="torch.utils._pytree._register_pytree_node is deprecated")
import json
from typing import List, Tuple, Dict
import anthropic
import numpy as np
import scipy.signal
import soundfile as sf
import torch
from tqdm import tqdm
import jax
import jax.numpy as jnp
from pydub import AudioSegment
from pydub.silence import detect_nonsilent
from whisper_jax import FlaxWhisperPipline
from pyannote.audio import Pipeline
from pyannote.core import Segment, Annotation
from docx import Document

class Swreporter:
    """A class for processing and analyzing audio files."""

    def __init__(self, audio_path: str, temp_output_dir: str, output_dir: str):
        self.audio_path = audio_path
        self.temp_output_dir = temp_output_dir
        self.output_dir = output_dir
        self.json_text = None

        # Claude settings
        self.claude_api_key = 'YOUR_API_KEY'
        # Consider using environment variables
        self.claude_client = anthropic.Anthropic(api_key=self.claude_api_key)
        self.prompt = None
        self.system_prompt = None
        # Audio processing settings
        self.jax_pipeline = FlaxWhisperPipline("openai/whisper-large-v3", dtype=jnp.bfloat16, batch_size=16)
        # First run, setup cache
        random_inputs = {
            "input_features": np.ones(
                (32, self.jax_pipeline.model.config.num_mel_bins, 2 * self.jax_pipeline.model.config.max_source_positions)
            )
        }
        self.jax_pipeline.forward(random_inputs, batch_size=32, return_timestamps=True)
        self.dia_pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1",
                                                     use_auth_token='YOUR_API_KEY').to(torch.device("cuda"))

    def _get_prompt(self) -> str:
        return f'''
                Write your own prompt
                逐字稿:
                {self.json_text}
                '''

    def _get_system_prompt(self) -> str:
        return '''Setup the Role charactor
                ''' 

    def split_on_silence(self, audio_path: str, min_silence_len: int = 500,
                         silence_thresh: int = -40, chunk_len: int = 600000) -> List[Tuple[int, int]]:
        """Split audio file based on silence."""
        audio = AudioSegment.from_file(audio_path, format='m4a')
        nonsilent_chunks = detect_nonsilent(audio, min_silence_len=min_silence_len, silence_thresh=silence_thresh)

        chunks = []
        start = 0
        for chunk_start, chunk_end in nonsilent_chunks:
            if chunk_end - start >= chunk_len:
                chunks.append((start, chunk_end))
                start = chunk_end

        if start < len(audio):
            chunks.append((start, len(audio)))

        return chunks

    def export_chunks(self, audio_path: str, chunks: List[Tuple[int, int]], output_folder: str):
        """Export audio chunks to separate files."""
        os.makedirs(output_folder, exist_ok=True)
        audio = AudioSegment.from_file(audio_path, format='m4a')
        for i, (start, end) in enumerate(chunks):
            chunk = audio[start:end]
            chunk.export(os.path.join(output_folder, f'chunk_{i}.wav'), format='wav')

    def segmenter(self, audio_path: str):
        """Segment the audio file."""
        chunks = self.split_on_silence(audio_path)
        self.export_chunks(audio_path, chunks, self.temp_output_dir)

    @staticmethod
    def equalizer(samples: np.ndarray, sample_rate: int) -> np.ndarray:
        """Apply equalization to audio samples."""
        b, a = scipy.signal.butter(4, 300 / (0.5 * sample_rate), btype='high')
        filtered_samples = scipy.signal.lfilter(b, a, samples)
        b, a = scipy.signal.butter(4, 3000 / (0.5 * sample_rate), btype='low')
        filtered_samples = scipy.signal.lfilter(b, a, filtered_samples)
        return filtered_samples

    def audio_preprocess(self):
        """Preprocess audio files."""
        data_list = [i for i in os.listdir(self.temp_output_dir) if i.endswith('.wav') and 'chunk' in i]
        for i, d in enumerate(data_list):
            audio = AudioSegment.from_file(f"{self.temp_output_dir}/{d}")
            audio = audio + 10  # Increase volume by 10dB
            samples = np.array(audio.get_array_of_samples())
            sample_rate = audio.frame_rate
            filtered_samples = self.equalizer(samples, sample_rate)
            enhanced_audio = audio._spawn(filtered_samples.astype(np.int16).tobytes())
            enhanced_audio.export(f"{self.temp_output_dir}/enhanced_segments_{i}.wav", format="wav")

        # Normalize
        data_list = [i for i in os.listdir(self.temp_output_dir) if i.endswith('.wav') and 'enhanced_segments' in i]
        for i, d in tqdm(enumerate(data_list), desc="enhancing audio..."):
            data, samplerate = sf.read(f'{self.temp_output_dir}/{d}')
            normalized_data = data / np.max(np.abs(data))
            sf.write(f'{self.temp_output_dir}/normalized_{i}.wav', normalized_data, samplerate)

    @staticmethod
    def intergration(speech: Dict, diarization: Annotation, label_mapping: Dict) -> List[str]:
        """Integrate speech recognition results with speaker diarization."""
        results = []
        for item in speech['segments']:
            start, end, text = item['start'], item['end'], item['text']
            speaker = diarization.crop(Segment(start, end)).argmax()
            line_short = f"{label_mapping[speaker]}:{text}"
            results.append(line_short)
        return results

    @staticmethod
    def list_to_word(sentence_list: List[str], output_doc: str):
        """Convert a list of sentences to a Word document."""
        doc = Document()
        for sentence in sentence_list:
            doc.add_paragraph(sentence)
        doc.save(output_doc)
        print(f'Transcribe done.')

    def save_report_to_word(self, report: str, filename: str):
        """
        :param report: the file need to be saved
        :param filename: file name
        """
        doc = Document()
        doc.add_paragraph(report)
        doc.save(filename)
        print(f"File saved as Word: {filename}")

    def transcribe(self):
        """Transcribe audio files."""
        soundfiles = [i for i in os.listdir(self.temp_output_dir) if i.endswith('.wav') and 'normalized' in i]
        speeches = []

        for d in tqdm(soundfiles, desc="Transcribing..."):
            text = self.jax_pipeline(f"{self.temp_output_dir}/{d}", return_timestamps=True, language='chinese')
            speeches.append(text)

        # Speaker separation
        data_list = [i for i in os.listdir(self.temp_output_dir) if i.endswith('.wav') and 'normalized' in i]
        results = []
        label_mapping = {'SPEAKER_00': 'a', 'SPEAKER_01': 'b', 'SPEAKER_02': 'c', None: 'unk'}

        for d, t in zip(data_list, speeches):
            diarization = self.dia_pipeline(f"{self.temp_output_dir}/{d}", min_speakers=2, max_speakers=2)
            speech = {'segments': [{'start': segment['timestamp'][0], 'end': segment['timestamp'][1], 'text': segment['text']} for segment in t['chunks']]}
            result = self.intergration(speech, diarization, label_mapping)
            results.extend(result)

        self.json_text = ','.join(results)
        self.list_to_word(results, os.path.join(self.output_dir, '逐字稿.docx'))

    def claude(self) -> str:
        """Generate report using Claude AI."""
        self.prompt = self._get_prompt()
        self.system_prompt = self._get_system_prompt()
        message = self.claude_client.messages.create(
            model="claude-3-5-sonnet-20240620",
            max_tokens=4096,
            temperature=0.7,
            system=self.system_prompt,
            messages=[{"role": "user", "content": self.prompt}]
        )
        return message.content[0].text

    @staticmethod
    def delete_specific_files_in_folder(folder_path: str, file_extension: str):
        """Delete specific files in a folder."""
        if not os.path.exists(folder_path):
            print(f"Folder {folder_path} does not exist")
            return

        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            if os.path.isfile(file_path) and filename.endswith(file_extension):
                try:
                    os.unlink(file_path)
                except Exception as e:
                    print(f"Wrong!, Delete manually. {file_path}: {e}")

    def run(self) -> str:
        """Run the entire process."""
        audio_data_name = next((i for i in os.listdir(self.audio_path) if i.endswith('.m4a')), None)
        if not audio_data_name:
            raise FileNotFoundError("You forget to put .m4a audio in folder.")

        self.segmenter(os.path.join(self.audio_path, audio_data_name))
        self.audio_preprocess()
        self.transcribe()
        self.delete_specific_files_in_folder(self.temp_output_dir, '.wav')
        report = self.claude()
        # Save as WORD file
        report_filename = os.path.join(self.output_dir, '報告.docx')
        self.save_report_to_word(report, report_filename)
